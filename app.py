import streamlit as st
import json
import io
from docx import Document
from docx.shared import Pt, Inches
from fpdf import FPDF
from PIL import Image

try:
    import openai
    st.write("OpenAI version:", getattr(openai, "__version__", "Desconocida"))
except ModuleNotFoundError:
    st.error("❌ La librería 'openai' no está instalada. Añádela en requirements.txt con:\n\nopenai\n")
    st.stop()

openai.api_key = st.secrets.get("OPENAI_API_KEY", "")
if not openai.api_key:
    st.error("❌ No se ha configurado la clave API de OpenAI. Añádela en Secrets como OPENAI_API_KEY.")
    st.stop()

st.set_page_config(page_title="Writing correction B1-B2", page_icon="✍️")
st.title("✍️ Writing correction - Bachillerato")
texto_alumno = st.text_area("📄 Pega aquí el writing del alumno:", height=200)

def evaluar_rubrica_con_gpt(texto_alumno):
    prompt = f"""
Eres un profesor de inglés. Evalúa el siguiente writing para un nivel B1-B2. Evalúa de forma realista y crítica. No asignes 0.5 a un criterio a menos que sea completamente correcto. La nota puede ser 0, 0.25 o 0.5. La suma total no debe superar 3 puntos según esta rúbrica: según esta rúbrica:

ADECUACIÓN (máximo 1.5 puntos)
- Cumplimiento de la tarea, registro y extensión (0.5)
- Variedad y organización de ideas (0.5)
- Cohesión y coherencia (0.5)

EXPRESIÓN (máximo 1.5 puntos)
- Gramática y estructuras (0.5)
- Vocabulario y riqueza léxica (0.5)
- Ortografía y puntuación (0.5)

Devolverás solo un JSON con los siguientes campos y ningún texto adicional:

{{
  "Adecuacion_Cumplimiento": 0.5,
  "Adecuacion_Variedad": 0.5,
  "Adecuacion_Cohesion": 0.5,
  "Expresion_Gramatica": 0.5,
  "Expresion_Vocabulario": 0.5,
  "Expresion_Ortografia": 0.5,
  "Justificaciones": {{
    "Cumplimiento": "detalles",
    "Variedad": "detalles",
    "Cohesion": "detalles",
    "Gramatica": "detalles",
    "Vocabulario": "detalles",
    "Ortografia": "detalles"
  }},
  "Errores_Detectados": "Lista completa de errores detectados, organizada por tipo de error (conjugación, artículos, concordancia, preposiciones, puntuación, estructuras confusas, capitalización). Cada sección debe empezar por ✅ seguido del número y tipo de error. A continuación, para cada error se presentará una tabla en texto plano con las columnas Error original, Corrección sugerida y Explicación, separadas por tabuladores (	), siguiendo el formato: 

✅ 1. Tipo de error
Error	Corrección	Explicación
Texto original	Texto corregido	Motivo de la corrección
...", siguiendo el formato: ✅ 1. Tipo de error.
Error	Corrección	Explicación
...
.",
  "Feedback": "Texto detallado explicando cómo mejorar en cada criterio de la rúbrica, basado en los errores concretos detectados.",
  "Writing_Reescrito": "Texto reescrito por la IA para obtener la nota máxima de 3/3."
}}

Texto a evaluar:
'''{texto_alumno}'''
"""
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
        max_tokens=1200,
    )
    return response.choices[0].message.content

if st.button("✅ Corregir"):
    if texto_alumno.strip():
        resultado_json = evaluar_rubrica_con_gpt(texto_alumno)
        try:
            start = resultado_json.find('{')
            end = resultado_json.rfind('}') + 1
            json_str = resultado_json[start:end]
            data = json.loads(json_str)
            criterios = {
                "Cumplimiento de la tarea": data.get("Adecuacion_Cumplimiento", 0),
                "Variedad y organización": data.get("Adecuacion_Variedad", 0),
                "Cohesión y coherencia": data.get("Adecuacion_Cohesion", 0),
                "Gramática": data.get("Expresion_Gramatica", 0),
                "Vocabulario": data.get("Expresion_Vocabulario", 0),
                "Ortografía y puntuación": data.get("Expresion_Ortografia", 0)
            }
            st.session_state['criterios'] = criterios
            st.session_state['data'] = data
            st.session_state['corregido'] = True
        except json.JSONDecodeError:
            st.error("❌ La respuesta de la IA no es un JSON válido.")

if 'criterios' in st.session_state and 'data' in st.session_state:
    criterios = st.session_state['criterios']
    data = st.session_state['data']
    if st.button("📥 Descargar informe en Word"):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        doc.add_heading("INFORME DE CORRECCIÓN - Writing", 0)
        try:
            doc.add_picture("logo_instituto.png", width=Inches(1.5))
        except:
            pass
        doc.add_heading("Writing original", level=1)
        p = doc.add_paragraph()
        errores = data.get("Errores_Detectados", "").lower()
        for word in texto_alumno.split():
            run = p.add_run(word + " ")
            if word in texto_alumno and word not in data.get("Writing_Reescrito", ""):
                if any(e in "grammar tenses verb conjugation" for e in errores):
                    run.font.highlight_color = 6  # Red highlight
                elif any(e in "cohesion coherence understand" for e in errores):
                    run.font.highlight_color = 7  # Yellow highlight
                elif any(e in "vocabulary lexis word choice" for e in errores):
                    run.font.highlight_color = 11  # Blue highlight
        doc.add_heading("Resultado de la rúbrica", level=1)
        for k, v in criterios.items():
            doc.add_paragraph(f"{k}: {v}/0.5")
        total_nota = sum(criterios.values())
        doc.add_paragraph(f"Nota total: {total_nota}/3")
        nota_sobre_10 = round((total_nota / 3) * 10, 2)
        doc.add_paragraph(f"Nota total: {nota_sobre_10}/10")
        doc.add_heading("Errores detectados", level=1)
        errores_texto = data.get("Errores_Detectados", "No disponible")
        for seccion in errores_texto.split("✅"):
            if seccion.strip():
                doc.add_heading(seccion.strip().splitlines()[0], level=2)
                tabla = doc.add_table(rows=1, cols=3)
                hdr_cells = tabla.rows[0].cells
                hdr_cells[0].text = 'Error'
                hdr_cells[1].text = 'Corrección'
                hdr_cells[2].text = 'Explicación'
                for linea in seccion.strip().splitlines():
                    if '	' in linea:
                        columnas = linea.split('	')
                        if len(columnas) == 3:
                            row_cells = tabla.add_row().cells
                            row_cells[0].text = columnas[0].strip()
                            row_cells[1].text = columnas[1].strip()
                            row_cells[2].text = columnas[2].strip()
        doc.add_heading("Feedback detallado", level=1)
        doc.add_paragraph(data.get("Feedback", "No disponible"))
        doc.add_heading("Writing reescrito para nota máxima (3/3)", level=1)
        doc.add_paragraph(data.get("Writing_Reescrito", "Texto reescrito no disponible"))
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("Descargar informe (Word)", data=buffer, file_name="informe_writing.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if st.button("📥 Descargar informe en PDF"):
        pdf = FPDF()
        pdf.add_page()
        try:
            pdf.image("logo_instituto.png", x=10, y=8, w=30)
        except:
            pass
        pdf.set_font("Arial", 'B', size=14)
        pdf.set_text_color(0, 0, 128)
        pdf.multi_cell(0, 10, "INFORME DE CORRECCIÓN - Writing\n\n")
        pdf.set_font("Arial", '', size=12)
        pdf.set_text_color(0, 0, 0)
        pdf.multi_cell(0, 10, f"Writing original:\n{texto_alumno}\n\n")
        pdf.multi_cell(0, 10, "Resultado de la rúbrica:\n")
        for k, v in criterios.items():
            pdf.multi_cell(0, 10, f"{k}: {v}/0.5")
        total_nota = sum(criterios.values())
        pdf.multi_cell(0, 10, f"Nota total: {total_nota}/3")
        nota_sobre_10 = round((total_nota / 3) * 10, 2)
        pdf.multi_cell(0, 10, f"Nota total: {nota_sobre_10}/10")
        errores_texto = data.get("Errores_Detectados", "No disponible")
        for seccion in errores_texto.split("✅"):
            if seccion.strip():
                pdf.set_font("Arial", 'B', 12)
                pdf.set_text_color(0, 0, 128)
                pdf.multi_cell(0, 10, seccion.strip().split("
")[0])
                pdf.set_font("Arial", '', 10)
                pdf.set_text_color(0, 0, 0)
                for linea in seccion.strip().split("
"):
                    if '	' in linea:
                        columnas = linea.split('	')
                        if len(columnas) == 3:
                            pdf.multi_cell(0, 5, f"Error: {columnas[0].strip()} | Corrección: {columnas[1].strip()} | Explicación: {columnas[2].strip()}")
        pdf.multi_cell(0, 10, f"\nFeedback detallado:\n{data.get('Feedback', 'No disponible')}\n")
        pdf.multi_cell(0, 10, "\nWriting reescrito para nota máxima (3/3):\n[Espacio para sugerencia del profesor]")
        buffer_pdf = io.BytesIO()
        pdf_bytes = pdf.output(dest='S').encode('latin1')
        buffer_pdf = io.BytesIO(pdf_bytes)
        buffer_pdf.seek(0)
        st.download_button("Descargar informe (PDF)", data=buffer_pdf, file_name="informe_writing.pdf", mime="application/pdf")
